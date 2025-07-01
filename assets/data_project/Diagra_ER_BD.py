from docx import Document
from docx.shared import Inches

# Crear documento de Word
doc = Document()

# Agregar título
doc.add_heading('Diagrama Entidad-Relación - RapiClean', level=1)

# Agregar descripción
doc.add_paragraph("Este diagrama muestra las principales entidades y sus relaciones en la base de datos RapiClean.")

# Insertar imagen del diagrama ER (si se hubiera generado antes) o agregar espacio para su edición manual
doc.add_paragraph("Diagrama de Entidad-Relación generado:")
doc.add_paragraph("[Aquí puedes insertar un diagrama creado con herramientas como DIA o dibujarlo manualmente.]")

# Guardar documento
file_path = "/mnt/data/Diagrama_ER_RapiClean.docx"
doc.save(file_path)

# Retornar el path del archivo generado
file_path
